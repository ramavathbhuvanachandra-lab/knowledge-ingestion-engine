from pathlib import Path
import json
import re

from docx import Document


# ============================================================
# CONFIG
# ============================================================

TEST_CASES = [
    {
        "domain": "iitj.ac.in",
        "organized_dir": Path("storage/organized_knowledge/iitj.ac.in"),
        "rag_dir": Path("storage/rag_knowledge/iitj.ac.in"),
        "expected_categories": 0,
        "expected_sections": 0,
        "reason": "IITJ test source contains only noise sections.",
    },
    {
        "domain": "gwpgc.ac.in",
        "organized_dir": Path("storage/organized_knowledge/gwpgc.ac.in"),
        "rag_dir": Path("storage/rag_knowledge/gwpgc.ac.in"),
        "expected_categories": 3,
        "expected_sections": 3,
        "reason": "GWPGC contains validated knowledge.",
    },
    {
        "domain": "home.iitd.ac.in",
        "organized_dir": Path("storage/organized_knowledge/home.iitd.ac.in"),
        "rag_dir": Path("storage/rag_knowledge/home.iitd.ac.in"),
        "expected_categories": 10,
        "expected_sections": 22,
        "reason": "IIT Delhi contains validated knowledge.",
    },
]


# ============================================================
# TEXT NORMALIZATION
# ============================================================

def normalize_text(text: str) -> str:
    """
    Normalize text for comparison between Markdown and DOCX.

    Markdown formatting differences such as:
        - bullet
        * bullet
        + bullet

    should not cause a false content-loss failure.
    """

    text = text.lower()

    # Remove common Markdown formatting.
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)

    # Remove Markdown heading markers.
    text = re.sub(
        r"^#+\s*",
        "",
        text,
        flags=re.MULTILINE,
    )

    # Remove Markdown unordered-list markers.
    text = re.sub(
        r"^\s*[-*+]\s+",
        "",
        text,
        flags=re.MULTILINE,
    )

    # Remove simple ordered-list markers.
    text = re.sub(
        r"^\s*\d+\.\s+",
        "",
        text,
        flags=re.MULTILINE,
    )

    # Normalize whitespace.
    text = " ".join(text.split())

    return text.strip()


# ============================================================
# FILE HELPERS
# ============================================================

def read_markdown(path: Path) -> str:
    return path.read_text(
        encoding="utf-8"
    )


def read_json(path: Path):
    return json.loads(
        path.read_text(
            encoding="utf-8"
        )
    )


def extract_docx_text(path: Path) -> str:
    """
    Extract text from:
    - DOCX paragraphs
    - DOCX tables
    """

    document = Document(path)

    parts = []

    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            parts.append(text)

    # --------------------------------------------------------
    # Tables
    # --------------------------------------------------------

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                text = cell.text.strip()

                if text:
                    parts.append(text)

    return "\n".join(parts)


def get_markdown_files(organized_dir: Path):
    """
    Actual organized knowledge structure:

        domain/
            category/
                *.md
    """

    return sorted(
        path
        for path in organized_dir.glob("*/*.md")
        if path.is_file()
    )


def get_docx_files(rag_dir: Path):
    """
    Actual RAG output structure:

        domain/
            category.docx
    """

    return sorted(
        path
        for path in rag_dir.glob("*.docx")
        if path.is_file()
    )


def get_organization_files(organized_dir: Path):
    """
    Organization metadata lives directly inside the
    domain directory.
    """

    return sorted(
        path
        for path in organized_dir.glob(
            "*.organization.json"
        )
        if path.is_file()
    )


def markdown_category(path: Path) -> str:
    """
    Category comes from the parent directory.

    Example:

        academics/newsletter.md

    -> academics
    """

    return path.parent.name


def docx_category(path: Path) -> str:
    """
    Category comes from the DOCX filename.

    Example:

        academics.docx

    -> academics
    """

    return path.stem


# ============================================================
# ORGANIZATION METADATA
# ============================================================

def load_organization_metadata(
    organized_dir: Path,
    domain: str,
):
    """
    Load the single organization JSON for the domain.
    """

    organization_files = get_organization_files(
        organized_dir
    )

    if len(organization_files) == 0:

        raise AssertionError(
            f"{domain}: no organization JSON found."
        )

    if len(organization_files) > 1:

        raise AssertionError(
            f"{domain}: multiple organization JSON files found: "
            f"{[path.name for path in organization_files]}"
        )

    organization_path = organization_files[0]

    metadata = read_json(
        organization_path
    )

    return organization_path, metadata


def validate_organization_metadata(
    metadata,
    domain: str,
    expected_categories: int,
    expected_sections: int,
):
    """
    Validate the structural contract produced by
    the organization phase.
    """

    required_fields = [
        "source_document",
        "domain",
        "document",
        "categories",
        "total_sections",
        "organized_sections",
        "written_files",
    ]

    # --------------------------------------------------------
    # Required fields
    # --------------------------------------------------------

    for field in required_fields:

        if field not in metadata:

            raise AssertionError(
                f"{domain}: organization metadata missing "
                f"required field: {field}"
            )

    # --------------------------------------------------------
    # Domain
    # --------------------------------------------------------

    if metadata["domain"] != domain:

        raise AssertionError(
            f"{domain}: organization metadata contains "
            f"wrong domain: {metadata['domain']}"
        )

    # --------------------------------------------------------
    # Categories
    # --------------------------------------------------------

    categories = metadata["categories"]

    if not isinstance(categories, dict):

        raise AssertionError(
            f"{domain}: 'categories' must be a dictionary."
        )

    active_categories = {
        category
        for category, count in categories.items()
        if count > 0
    }

    if len(active_categories) != expected_categories:

        raise AssertionError(
            f"{domain}: organization metadata contains "
            f"{len(active_categories)} active categories, "
            f"expected {expected_categories}"
        )

    # --------------------------------------------------------
    # Section totals
    # --------------------------------------------------------

    total_sections = metadata["total_sections"]

    organized_sections = metadata[
        "organized_sections"
    ]

    if total_sections != expected_sections:

        raise AssertionError(
            f"{domain}: organization metadata reports "
            f"{total_sections} total sections, "
            f"expected {expected_sections}"
        )

    if organized_sections != expected_sections:

        raise AssertionError(
            f"{domain}: organization metadata reports "
            f"{organized_sections} organized sections, "
            f"expected {expected_sections}"
        )

    # --------------------------------------------------------
    # Category section sum
    # --------------------------------------------------------

    category_section_total = sum(
        count
        for count in categories.values()
    )

    if category_section_total != total_sections:

        raise AssertionError(
            f"{domain}: category section counts sum to "
            f"{category_section_total}, but total_sections "
            f"is {total_sections}"
        )

    return categories


# ============================================================
# MARKDOWN STRUCTURE
# ============================================================

def validate_markdown_mapping(
    markdown_files,
    categories,
    domain: str,
):
    """
    Validate that organized Markdown files match the
    active categories from organization.json.
    """

    markdown_categories = {
        markdown_category(path)
        for path in markdown_files
    }

    expected_categories = {
        category
        for category, count in categories.items()
        if count > 0
    }

    # --------------------------------------------------------
    # Missing categories
    # --------------------------------------------------------

    missing_categories = (
        expected_categories
        - markdown_categories
    )

    if missing_categories:

        raise AssertionError(
            f"{domain}: missing organized Markdown "
            f"categories: {sorted(missing_categories)}"
        )

    # --------------------------------------------------------
    # Unexpected categories
    # --------------------------------------------------------

    unexpected_categories = (
        markdown_categories
        - expected_categories
    )

    if unexpected_categories:

        raise AssertionError(
            f"{domain}: unexpected organized Markdown "
            f"categories: {sorted(unexpected_categories)}"
        )

    # --------------------------------------------------------
    # One Markdown file per active category
    # --------------------------------------------------------

    for category in sorted(
        expected_categories
    ):

        category_files = [
            path
            for path in markdown_files
            if markdown_category(path) == category
        ]

        if len(category_files) != 1:

            raise AssertionError(
                f"{domain}: category '{category}' has "
                f"{len(category_files)} Markdown files; "
                f"expected exactly 1."
            )


# ============================================================
# SOURCE CONTENT EXTRACTION
# ============================================================

def is_metadata_line(line: str) -> bool:
    """
    Identify internal Markdown metadata that does not need
    to survive into the final RAG DOCX.
    """

    normalized = normalize_text(line)

    metadata_prefixes = (
        "source document:",
        "source:",
        "domain:",
        "document:",
        "category:",
        "generated:",
        "generated at:",
        "metadata:",
    )

    return normalized.startswith(
        metadata_prefixes
    )


def extract_knowledge_lines(markdown_text: str):
    """
    Extract meaningful knowledge content.

    Excludes:
    - Markdown headings
    - code fences
    - internal metadata
    - tiny formatting fragments

    Markdown list markers are NOT considered knowledge.
    They are removed during normalization.
    """

    lines = []

    for raw_line in markdown_text.splitlines():

        line = raw_line.strip()

        if not line:
            continue

        # ----------------------------------------------------
        # Markdown headings
        # ----------------------------------------------------

        if line.startswith("#"):
            continue

        # ----------------------------------------------------
        # Code fences
        # ----------------------------------------------------

        if line.startswith("```"):
            continue

        # ----------------------------------------------------
        # Internal metadata
        # ----------------------------------------------------

        if is_metadata_line(line):
            continue

        # ----------------------------------------------------
        # Normalize
        # ----------------------------------------------------

        normalized = normalize_text(line)

        if len(normalized) < 5:
            continue

        lines.append(normalized)

    return lines


# ============================================================
# CONTENT VALIDATION
# ============================================================

def validate_content_preservation(
    markdown_text: str,
    docx_text: str,
    domain: str,
    category: str,
):
    """
    Validate that meaningful knowledge from Markdown survives
    into the DOCX.

    We do NOT require:
        Markdown == DOCX

    We DO require meaningful source knowledge to survive.

    Formatting differences such as Markdown bullets are ignored.
    """

    normalized_docx = normalize_text(
        docx_text
    )

    source_lines = extract_knowledge_lines(
        markdown_text
    )

    if not source_lines:
        return

    missing = []

    for line in source_lines:

        if line not in normalized_docx:

            missing.append(line)

    if missing:

        preview = missing[:5]

        raise AssertionError(
            f"{domain}: significant knowledge missing "
            f"from {category}.docx. "
            f"Examples: {preview}"
        )


# ============================================================
# VALIDATE ONE DOMAIN
# ============================================================

def validate_domain(case):

    domain = case["domain"]

    organized_dir = case[
        "organized_dir"
    ]

    rag_dir = case[
        "rag_dir"
    ]

    expected_categories = case[
        "expected_categories"
    ]

    expected_sections = case[
        "expected_sections"
    ]

    print("=" * 100)

    print(
        f" DOMAIN    : {domain}"
    )

    print(
        f" ORGANIZED : {organized_dir}"
    )

    print(
        f" RAG       : {rag_dir}"
    )

    print(
        f" Reason    : {case['reason']}"
    )

    print()

    # ========================================================
    # DIRECTORY CHECK
    # ========================================================

    if not organized_dir.exists():

        raise AssertionError(
            f"{domain}: organized directory does not exist: "
            f"{organized_dir}"
        )

    if not rag_dir.exists():

        raise AssertionError(
            f"{domain}: RAG directory does not exist: "
            f"{rag_dir}"
        )

    # ========================================================
    # DISCOVERY
    # ========================================================

    markdown_files = get_markdown_files(
        organized_dir
    )

    docx_files = get_docx_files(
        rag_dir
    )

    print(
        f" Markdown files : {len(markdown_files)}"
    )

    print(
        f" DOCX files     : {len(docx_files)}"
    )

    # ========================================================
    # ZERO-KNOWLEDGE CASE
    # ========================================================

    if expected_categories == 0:

        if markdown_files:

            raise AssertionError(
                f"{domain}: expected zero organized "
                f"Markdown files, found "
                f"{len(markdown_files)}"
            )

        if docx_files:

            raise AssertionError(
                f"{domain}: unexpected DOCX files "
                f"generated: "
                f"{[path.name for path in docx_files]}"
            )

        print(
            " Categories     : 0"
        )

        print(
            " Sections       : 0"
        )

        print()

        print(
            " PASS — no knowledge correctly produced no DOCX."
        )

        print()

        return

    # ========================================================
    # ORGANIZATION JSON
    # ========================================================

    organization_path, metadata = (
        load_organization_metadata(
            organized_dir,
            domain,
        )
    )

    categories = validate_organization_metadata(
        metadata,
        domain,
        expected_categories,
        expected_sections,
    )

    active_categories = {
        category
        for category, count
        in categories.items()
        if count > 0
    }

    print(
        f" Categories     : "
        f"{len(active_categories)}"
    )

    print(
        f" Sections       : "
        f"{metadata['organized_sections']}"
    )

    print(
        f" Metadata       : "
        f"{organization_path.name}"
    )

    # ========================================================
    # MARKDOWN STRUCTURE
    # ========================================================

    validate_markdown_mapping(
        markdown_files,
        categories,
        domain,
    )

    # ========================================================
    # DOCX COUNT
    # ========================================================

    if len(docx_files) != expected_categories:

        raise AssertionError(
            f"{domain}: expected "
            f"{expected_categories} DOCX files, "
            f"found {len(docx_files)}"
        )

    # ========================================================
    # DOCX CATEGORY MAPPING
    # ========================================================

    docx_categories = {
        docx_category(path)
        for path in docx_files
    }

    missing_docx_categories = (
        active_categories
        - docx_categories
    )

    unexpected_docx_categories = (
        docx_categories
        - active_categories
    )

    if missing_docx_categories:

        raise AssertionError(
            f"{domain}: missing DOCX categories: "
            f"{sorted(missing_docx_categories)}"
        )

    if unexpected_docx_categories:

        raise AssertionError(
            f"{domain}: unexpected DOCX categories: "
            f"{sorted(unexpected_docx_categories)}"
        )

    # ========================================================
    # DUPLICATE DOCX CHECK
    # ========================================================

    docx_category_counts = {}

    for path in docx_files:

        category = docx_category(path)

        docx_category_counts[category] = (
            docx_category_counts.get(
                category,
                0
            ) + 1
        )

    duplicates = [
        category
        for category, count
        in docx_category_counts.items()
        if count > 1
    ]

    if duplicates:

        raise AssertionError(
            f"{domain}: duplicate DOCX categories "
            f"detected: {duplicates}"
        )

    # ========================================================
    # CATEGORY VALIDATION
    # ========================================================

    for category in sorted(
        active_categories
    ):

        print(
            f" VALIDATING : {category}"
        )

        # ----------------------------------------------------
        # Source Markdown
        # ----------------------------------------------------

        category_markdown_files = [
            path
            for path in markdown_files
            if markdown_category(path)
            == category
        ]

        if len(category_markdown_files) != 1:

            raise AssertionError(
                f"{domain}: category '{category}' "
                f"must contain exactly one Markdown file."
            )

        markdown_path = (
            category_markdown_files[0]
        )

        # ----------------------------------------------------
        # Expected DOCX
        # ----------------------------------------------------

        expected_docx = (
            rag_dir
            / f"{category}.docx"
        )

        if not expected_docx.exists():

            raise AssertionError(
                f"{domain}: missing expected DOCX: "
                f"{expected_docx}"
            )

        # ----------------------------------------------------
        # DOCX must contain text
        # ----------------------------------------------------

        docx_text = extract_docx_text(
            expected_docx
        )

        if not docx_text.strip():

            raise AssertionError(
                f"{domain}: DOCX is empty: "
                f"{expected_docx.name}"
            )

        # ----------------------------------------------------
        # Knowledge preservation
        # ----------------------------------------------------

        markdown_text = read_markdown(
            markdown_path
        )

        validate_content_preservation(
            markdown_text,
            docx_text,
            domain,
            category,
        )

        print(
            "   PASS"
        )

    # ========================================================
    # DOMAIN PASS
    # ========================================================

    print()

    print(
        f" PASS — {domain}"
    )

    print()


# ============================================================
# MAIN
# ============================================================

def main():

    passed = 0

    for case in TEST_CASES:

        validate_domain(case)

        passed += 1

    print("=" * 100)

    print(
        f" Domains passed: "
        f"{passed} / {len(TEST_CASES)}"
    )

    print()

    print(
        "# PHASE 8.6 RAG DOCUMENT VALIDATION: PASS"
    )


# ============================================================
# ENTRY POINT
# ============================================================

if __name__ == "__main__":
    main()