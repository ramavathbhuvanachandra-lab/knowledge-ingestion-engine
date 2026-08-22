"""
IIT Jodhpur — Command 5 Traceable DOCX Exporter V2
==================================================

Source of truth:
    storage_v2/command5/iitj_final/*.md

Output:
    storage_v2/command5/iitj_final_docs_v2/*.docx

Design goals
------------
1. Human-readable and editable.
2. Preserve Command 5 Markdown as canonical source-of-truth.
3. Clean metadata header instead of raw YAML/Markdown.
4. Content is rendered as a proper information document.
5. Duplicate visible knowledge inside the same document is collapsed when
   the normalized heading + normalized content are identical.
6. All original unit IDs remain traceable in a dedicated source section.
7. Source documents and source URLs remain visible and clickable.
8. Raw Markdown syntax is removed from visible DOCX text where formatting can
   be represented natively.
9. No semantic rewriting, summarization, invention, or deletion of unique
   source knowledge.
10. Duplicate provenance is retained even when duplicate visible content is
    collapsed.

This exporter changes presentation only. It does NOT modify Command 5
Markdown or its organization plan/manifest.
"""

from __future__ import annotations

import json
import re
from collections import defaultdict
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[1]

SOURCE_ROOT = (
    PROJECT_ROOT
    / "storage_v2"
    / "command5"
    / "iitj_final"
)

OUTPUT_ROOT = (
    PROJECT_ROOT
    / "storage_v2"
    / "command5"
    / "iitj_final_docs_v2"
)

PLAN_PATH = SOURCE_ROOT / "iitj_organization_plan.json"
MANIFEST_PATH = SOURCE_ROOT / "iitj_organization_manifest.json"


# ---------------------------------------------------------------------------
# Formatting helpers.
# ---------------------------------------------------------------------------

def set_font(
    run,
    name: str = "Aptos",
    size: float = 10.5,
    bold: bool | None = None,
) -> None:
    run.font.name = name
    run.font.size = Pt(size)

    if bold is not None:
        run.bold = bold

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)

    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run.append(rpr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_rule(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.space_before = Pt(5)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7C4D6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_font(run, size=9)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size in (
        ("Title", 22),
        ("Heading 1", 16),
        ("Heading 2", 13),
        ("Heading 3", 11),
    ):
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = footer_p.add_run("IIT Jodhpur • Command 5")
    set_font(run, size=8.5)

    footer_p.add_run("  •  Page ")
    add_page_number(footer_p)


# ---------------------------------------------------------------------------
# Metadata / Markdown parsing.
# ---------------------------------------------------------------------------

def parse_front_matter(text: str) -> tuple[dict, str]:
    text = text.lstrip()

    if not text.startswith("---"):
        return {}, text

    pieces = text.split("---", 2)
    if len(pieces) != 3:
        return {}, text

    meta_raw = pieces[1]
    body = pieces[2].lstrip()

    meta = {}
    current_list_key = None

    for raw in meta_raw.splitlines():
        line = raw.rstrip()

        if not line.strip():
            continue

        bullet = re.match(r"^\s*-\s+(.*)$", line)
        if bullet and current_list_key:
            meta.setdefault(current_list_key, []).append(
                bullet.group(1).strip()
            )
            continue

        match = re.match(
            r"^([A-Za-z0-9_]+):\s*(.*)$",
            line,
        )
        if not match:
            continue

        key = match.group(1)
        value = match.group(2).strip()

        if value == "":
            meta[key] = []
            current_list_key = key
        else:
            meta[key] = value
            current_list_key = None

    return meta, body


def clean_inline_markdown(text: str) -> str:
    text = text.replace("\r", "")

    # Markdown links -> visible label, URL remains handled by traceability.
    text = re.sub(
        r"\[([^\]]+)\]\(([^)]+)\)",
        r"\1",
        text,
    )

    # Bold/italic/code markers.
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"_(.*?)_", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)

    # Normalize common separator artifacts without changing words.
    text = text.replace("||", " | ")
    text = re.sub(r"\s{2,}", " ", text)

    return text.strip()


def parse_body(body: str) -> list[tuple[str, str]]:
    blocks = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return

        value = " ".join(
            line.strip()
            for line in paragraph_lines
        ).strip()

        if value:
            blocks.append(
                ("paragraph", clean_inline_markdown(value))
            )

        paragraph_lines.clear()

    for raw in body.splitlines():
        line = raw.rstrip()

        if not line.strip():
            flush_paragraph()
            continue

        heading = re.match(
            r"^(#{1,6})\s+(.*)$",
            line,
        )

        if heading:
            flush_paragraph()
            blocks.append(
                (
                    f"heading_{len(heading.group(1))}",
                    clean_inline_markdown(
                        heading.group(2).strip()
                    ),
                )
            )
            continue

        bullet = re.match(
            r"^\s*-\s+(.*)$",
            line,
        )

        if bullet:
            flush_paragraph()
            blocks.append(
                (
                    "bullet",
                    clean_inline_markdown(
                        bullet.group(1).strip()
                    ),
                )
            )
            continue

        paragraph_lines.append(line)

    flush_paragraph()
    return blocks


# ---------------------------------------------------------------------------
# Semantic/presentation deduplication.
# ---------------------------------------------------------------------------

def normalize_for_duplicate(text: str) -> str:
    text = clean_inline_markdown(text).lower()
    text = re.sub(r"[\u2010-\u2015]", "-", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def dedupe_blocks(
    blocks: list[tuple[str, str]],
) -> tuple[list[tuple[str, str, int]], int]:
    """
    Collapse repeated content blocks.

    IMPORTANT:
    - Only exact normalized duplicates are collapsed.
    - Near-duplicates are NOT merged.
    - No unique information is removed.
    - Returns duplicate count for QA.
    """
    seen = {}
    result = []
    duplicate_count = 0

    for kind, value in blocks:
        normalized_value = normalize_for_duplicate(value)

        # Exact duplicate visible text is collapsed even if one occurrence
        # arrived as a heading and another as a paragraph. This handles common
        # scrape artifacts such as:
        #   "Associated Faculty Members"
        #   Associated Faculty Members
        # while preserving genuinely different wording.
        if kind == "bullet":
            key = ("content", normalized_value)
        elif kind.startswith("heading_"):
            key = ("content", normalized_value)
        else:
            key = ("content", normalized_value)

        if not key[1]:
            continue

        if key in seen:
            seen[key]["count"] += 1
            duplicate_count += 1
            continue

        entry = {
            "kind": kind,
            "value": value,
            "count": 1,
        }
        seen[key] = entry
        result.append(entry)

    return [
        (
            item["kind"],
            item["value"],
            item["count"],
        )
        for item in result
    ], duplicate_count


# ---------------------------------------------------------------------------
# Metadata / provenance presentation.
# ---------------------------------------------------------------------------

def add_metadata_header(
    document: Document,
    metadata: dict,
    source_urls: list[str],
) -> None:
    title = str(
        metadata.get("title", "")
    ).strip()

    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        title or "IIT Jodhpur Knowledge Document"
    )
    set_font(run, size=22)
    run.bold = True

    subtitle = document.add_paragraph()
    r = subtitle.add_run(
        "IIT Jodhpur • Command 5 Knowledge Document"
    )
    set_font(r, size=10.5)
    r.bold = True

    add_rule(document)

    table = document.add_table(
        rows=0,
        cols=2,
    )
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    fields = [
        ("Document ID", metadata.get("id", "")),
        ("Scope", metadata.get("scope_type", "")),
        ("Scope ID", metadata.get("scope_id", "")),
        ("Topic", metadata.get("topic", "")),
        (
            "Organization Version",
            metadata.get(
                "organization_version",
                "",
            ),
        ),
        (
            "Knowledge Units",
            str(
                len(
                    metadata.get(
                        "unit_ids",
                        [],
                    )
                )
            ),
        ),
        (
            "Source",
            "IIT Jodhpur Official Website",
        ),
    ]

    for label, value in fields:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        left = cells[0].paragraphs[0].add_run(label)
        set_font(left, size=9.5, bold=True)

        right = cells[1].paragraphs[0].add_run(
            str(value)
        )
        set_font(right, size=9.5)

    if source_urls:
        cells = table.add_row().cells
        left = cells[0].paragraphs[0].add_run(
            "Primary Source URL(s)"
        )
        set_font(left, size=9.5, bold=True)

        p = cells[1].paragraphs[0]
        for i, url in enumerate(source_urls):
            add_hyperlink(
                p,
                url,
                url,
            )
            if i < len(source_urls) - 1:
                p.add_run("\n")

    document.add_paragraph()
    add_rule(document)


def add_traceability(
    document: Document,
    rows: list[dict],
) -> None:
    if not rows:
        return

    p = document.add_paragraph(style="Heading 1")
    r = p.add_run("Source Traceability")
    set_font(r, size=16)

    intro = document.add_paragraph()
    r = intro.add_run(
        "Every knowledge unit remains linked to its source document "
        "and original source URL."
    )
    set_font(r, size=9.5)

    table = document.add_table(
        rows=1,
        cols=3,
    )
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = (
        "Knowledge Unit ID",
        "Source Document",
        "Source URL",
    )

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(header)
        set_font(run, size=9.2, bold=True)

    for row in rows:
        cells = table.add_row().cells

        unit_run = cells[0].paragraphs[0].add_run(
            row.get("unit_id", "")
        )
        set_font(unit_run, size=8.7)

        file_run = cells[1].paragraphs[0].add_run(
            row.get("source_file", "")
        )
        set_font(file_run, size=8.7)

        p = cells[2].paragraphs[0]
        url = row.get("source_url", "")
        if url:
            add_hyperlink(
                p,
                url,
                url,
            )

    document.add_paragraph()


# ---------------------------------------------------------------------------
# Content presentation.
# ---------------------------------------------------------------------------

def add_heading(
    document: Document,
    level: int,
    text: str,
) -> None:
    if level == 1:
        style = "Heading 1"
        size = 16
    elif level == 2:
        style = "Heading 2"
        size = 13
    else:
        style = "Heading 3"
        size = 11

    p = document.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True

    run = p.add_run(text)
    set_font(
        run,
        size=size,
        bold=True,
    )


def add_bullet(
    document: Document,
    text: str,
) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)

    run = p.add_run(text)
    set_font(run, size=10.2)


def add_paragraph(
    document: Document,
    text: str,
) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)

    run = p.add_run(text)
    set_font(run, size=10.5)


def render_content(
    document: Document,
    blocks: list[tuple[str, str, int]],
) -> dict:
    visible_blocks = 0
    collapsed_duplicates = 0
    first_heading_1_seen = False

    for kind, value, count in blocks:
        if count > 1:
            collapsed_duplicates += count - 1

        if kind == "heading_1":
            if first_heading_1_seen:
                # The organizer title is already in the metadata header.
                add_heading(
                    document,
                    1,
                    value,
                )
            else:
                first_heading_1_seen = True
            visible_blocks += 1
            continue

        if kind == "heading_2":
            add_heading(
                document,
                2,
                value,
            )
            visible_blocks += 1
            continue

        if kind.startswith("heading_"):
            add_heading(
                document,
                3,
                value,
            )
            visible_blocks += 1
            continue

        if kind == "bullet":
            add_bullet(
                document,
                value,
            )
            visible_blocks += 1
            continue

        add_paragraph(
            document,
            value,
        )
        visible_blocks += 1

    return {
        "visible_blocks": visible_blocks,
        "collapsed_duplicates": collapsed_duplicates,
    }


# ---------------------------------------------------------------------------
# File-level conversion.
# ---------------------------------------------------------------------------

def load_indexes() -> tuple[dict, dict]:
    plan = json.loads(
        PLAN_PATH.read_text(
            encoding="utf-8",
        )
    )

    manifest = json.loads(
        MANIFEST_PATH.read_text(
            encoding="utf-8",
        )
    )

    return plan, manifest


def build_manifest_index(manifest: dict) -> dict[str, dict]:
    return {
        row["unit_id"]: row
        for row in manifest.get(
            "assignments",
            [],
        )
    }


def build_trace_rows(
    unit_ids: Iterable[str],
    manifest_by_unit: dict[str, dict],
) -> list[dict]:
    rows = []

    for unit_id in unit_ids:
        row = manifest_by_unit.get(unit_id)

        if not row:
            rows.append(
                {
                    "unit_id": unit_id,
                    "source_file": "",
                    "source_url": "",
                }
            )
            continue

        rows.append(row)

    return rows


def convert_document(
    markdown_path: Path,
    document_plan: dict,
    manifest_by_unit: dict[str, dict],
    output_path: Path,
) -> dict:
    text = markdown_path.read_text(
        encoding="utf-8",
        errors="replace",
    )

    front_matter, body = parse_front_matter(text)

    metadata = dict(front_matter)

    for key in (
        "id",
        "title",
        "scope_type",
        "scope_id",
        "topic",
    ):
        if document_plan.get(key) is not None:
            metadata[key] = document_plan[key]

    metadata["unit_ids"] = document_plan.get(
        "unit_ids",
        metadata.get(
            "unit_ids",
            [],
        ),
    )

    metadata["organization_version"] = (
        metadata.get(
            "organization_version",
            "",
        )
        or document_plan.get(
            "organization_version",
            "",
        )
    )

    trace_rows = build_trace_rows(
        metadata["unit_ids"],
        manifest_by_unit,
    )

    source_urls = sorted(
        {
            row.get("source_url", "")
            for row in trace_rows
            if row.get("source_url")
        }
    )

    blocks = parse_body(body)

    # Do not render the Markdown's own raw traceability section as body text.
    trace_start = None
    for index, (kind, value) in enumerate(blocks):
        if (
            kind == "heading_2"
            and normalize_for_duplicate(value)
            == "source traceability"
        ):
            trace_start = index
            break

    if trace_start is not None:
        blocks = blocks[:trace_start]

    deduped_blocks, duplicate_count = dedupe_blocks(
        blocks
    )

    document = Document()
    configure_document(document)

    add_metadata_header(
        document,
        metadata,
        source_urls,
    )

    content_stats = render_content(
        document,
        deduped_blocks,
    )

    # Add a compact presentation note only when duplicates actually existed.
    if duplicate_count:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(6)

        run = p.add_run(
            f"Presentation note: {duplicate_count} exact duplicate "
            "content occurrences were collapsed for readability. "
            "All original unit IDs remain preserved in Source Traceability."
        )
        set_font(run, size=8.8)
        run.italic = True

    add_rule(document)

    add_traceability(
        document,
        trace_rows,
    )

    document.core_properties.title = str(
        metadata.get(
            "title",
            markdown_path.stem,
        )
    )
    document.core_properties.subject = (
        "IIT Jodhpur Command 5 human-readable knowledge document"
    )
    document.core_properties.keywords = (
        "IIT Jodhpur, Command 5, knowledge, traceability"
    )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(output_path)

    return {
        "source_markdown": str(
            markdown_path.relative_to(SOURCE_ROOT)
        ),
        "output_docx": str(
            output_path.relative_to(OUTPUT_ROOT)
        ),
        "source_unit_count": len(
            metadata["unit_ids"]
        ),
        "source_block_count": len(blocks),
        "visible_block_count": content_stats[
            "visible_blocks"
        ],
        "collapsed_duplicate_occurrences": (
            duplicate_count
        ),
        "traceability_row_count": len(
            trace_rows
        ),
        "source_url_count": len(source_urls),
    }


def main() -> None:
    if not SOURCE_ROOT.exists():
        raise FileNotFoundError(
            f"Missing source tree: {SOURCE_ROOT}"
        )

    if not PLAN_PATH.exists():
        raise FileNotFoundError(
            f"Missing organization plan: {PLAN_PATH}"
        )

    if not MANIFEST_PATH.exists():
        raise FileNotFoundError(
            f"Missing organization manifest: {MANIFEST_PATH}"
        )

    plan, manifest = load_indexes()
    manifest_by_unit = build_manifest_index(
        manifest
    )

    plan_by_path = {
        item["path"]: item
        for item in plan.get(
            "documents",
            [],
        )
    }

    markdown_files = sorted(
        SOURCE_ROOT.rglob("*.md")
    )

    excluded = {
        "iitj_semantic_audit.md",
        "iitj_buffer_inspection.md",
        "iitj_semantic_audit.json",
        "iitj_buffer_inspection.json",
    }

    markdown_files = [
        path
        for path in markdown_files
        if path.name not in excluded
    ]

    if not markdown_files:
        raise RuntimeError(
            "No Command 5 Markdown documents found."
        )

    # Fresh generated folder. This does not touch the canonical Markdown.
    if OUTPUT_ROOT.exists():
        shutil.rmtree(OUTPUT_ROOT)

    OUTPUT_ROOT.mkdir(
        parents=True,
        exist_ok=True,
    )

    reports = []
    total_duplicates = 0
    total_units = 0
    total_trace_rows = 0

    for markdown_path in markdown_files:
        relative = markdown_path.relative_to(
            SOURCE_ROOT
        )

        document_plan = plan_by_path.get(
            str(relative)
        )

        if not document_plan:
            raise ValueError(
                "Markdown document missing from organization plan: "
                + str(relative)
            )

        output_path = (
            OUTPUT_ROOT
            / relative.parent
            / f"{relative.stem}.docx"
        )

        report = convert_document(
            markdown_path,
            document_plan,
            manifest_by_unit,
            output_path,
        )

        reports.append(report)
        total_duplicates += report[
            "collapsed_duplicate_occurrences"
        ]
        total_units += report[
            "source_unit_count"
        ]
        total_trace_rows += report[
            "traceability_row_count"
        ]

    report_path = (
        OUTPUT_ROOT
        / "docx_export_report.json"
    )

    report_data = {
        "export_version": "iitj_command5_docs_v2",
        "source_of_truth": str(
            SOURCE_ROOT
        ),
        "output_root": str(
            OUTPUT_ROOT
        ),
        "document_count": len(reports),
        "total_source_unit_references": total_units,
        "total_traceability_rows": total_trace_rows,
        "total_collapsed_duplicate_occurrences": (
            total_duplicates
        ),
        "documents": reports,
    }

    report_path.write_text(
        json.dumps(
            report_data,
            indent=2,
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    actual_docx = sorted(
        OUTPUT_ROOT.rglob("*.docx")
    )

    if len(actual_docx) != len(
        markdown_files
    ):
        raise RuntimeError(
            "DOCX count mismatch: "
            f"expected={len(markdown_files)} "
            f"actual={len(actual_docx)}"
        )

    print("=" * 100)
    print("IIT JODHPUR — COMMAND 5 TRACEABLE DOCX V2")
    print("=" * 100)
    print()
    print("Markdown source documents:", len(markdown_files))
    print("DOCX documents created:", len(actual_docx))
    print(
        "Exact duplicate occurrences collapsed:",
        total_duplicates,
    )
    print(
        "Traceability rows preserved:",
        total_trace_rows,
    )
    print()
    print("Output:", OUTPUT_ROOT)
    print("Report:", report_path)
    print()
    print(
        "Source of truth remains:",
        SOURCE_ROOT,
    )


if __name__ == "__main__":
    main()