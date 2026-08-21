"""
IIT Jodhpur — Command 5 Markdown -> DOCX V1

Converts every generated Markdown document under:
    storage_v2/command5/iitj_final/

into a matching DOCX tree under:
    storage_v2/command5/iitj_final_docs/

The Markdown remains the canonical intermediate output.
The DOCX layer is a presentation/export artifact and never rewrites source
content.

Supported Markdown:
- YAML-style metadata blocks are skipped from visible body
- #, ##, ### headings
- unordered list items
- plain paragraphs
- source-traceability section
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[1]

SOURCE_ROOT = (
    PROJECT_ROOT
    / "storage_v2"
    / "command5"
    / "iitj_final"
)

DOCX_ROOT = (
    PROJECT_ROOT
    / "storage_v2"
    / "command5"
    / "iitj_final_docs"
)


def set_run_font(run, name: str = "Aptos", size: float = 10.5) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(
        qn("w:ascii"),
        name,
    )
    run._element.rPr.rFonts.set(
        qn("w:hAnsi"),
        name,
    )


def strip_front_matter(text: str) -> str:
    text = text.lstrip()

    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) == 3:
            return parts[2].lstrip()

    return text


def parse_markdown(text: str) -> list[tuple[str, str]]:
    text = strip_front_matter(text)

    blocks = []
    paragraph_lines = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            value = " ".join(
                line.strip()
                for line in paragraph_lines
            ).strip()

            if value:
                blocks.append(
                    ("paragraph", value)
                )

            paragraph_lines.clear()

    for raw in text.splitlines():
        line = raw.rstrip()

        if not line.strip():
            flush_paragraph()
            continue

        heading = re.match(
            r"^(#{1,6})\s+(.*)$",
            line,
        )

        if heading:
            flush_paragraph()
            level = len(
                heading.group(1)
            )
            blocks.append(
                (f"h{level}", heading.group(2).strip())
            )
            continue

        bullet = re.match(
            r"^\s*-\s+(.*)$",
            line,
        )

        if bullet:
            flush_paragraph()
            blocks.append(
                ("bullet", bullet.group(1).strip())
            )
            continue

        paragraph_lines.append(line)

    flush_paragraph()
    return blocks


def create_docx(markdown_path: Path, output_path: Path) -> None:
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Base font.
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(
        qn("w:ascii"),
        "Aptos",
    )
    normal._element.rPr.rFonts.set(
        qn("w:hAnsi"),
        "Aptos",
    )

    blocks = parse_markdown(
        markdown_path.read_text(
            encoding="utf-8",
            errors="replace",
        )
    )

    first_heading = True

    for kind, value in blocks:
        if kind == "h1":
            p = document.add_paragraph(
                style="Title"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(
                run,
                size=20,
            )
            first_heading = False

        elif kind == "h2":
            p = document.add_paragraph(
                style="Heading 1"
            )
            run = p.add_run(value)
            set_run_font(
                run,
                size=15,
            )

        elif kind == "h3":
            p = document.add_paragraph(
                style="Heading 2"
            )
            run = p.add_run(value)
            set_run_font(
                run,
                size=12,
            )

        elif kind == "h4":
            p = document.add_paragraph(
                style="Heading 3"
            )
            run = p.add_run(value)
            set_run_font(
                run,
                size=11,
            )

        elif kind == "bullet":
            p = document.add_paragraph(
                style="List Bullet"
            )
            run = p.add_run(value)
            set_run_font(run)

        else:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(value)
            set_run_font(run)

    document.core_properties.title = markdown_path.stem
    document.core_properties.subject = (
        "IIT Jodhpur Command 5 organized knowledge"
    )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )
    document.save(output_path)


def main() -> None:
    if not SOURCE_ROOT.exists():
        raise FileNotFoundError(
            f"Missing source tree: {SOURCE_ROOT}"
        )

    markdown_files = sorted(
        SOURCE_ROOT.rglob("*.md")
    )

    # Avoid converting audit reports if this script is rerun.
    markdown_files = [
        path
        for path in markdown_files
        if path.name
        not in {
            "iitj_semantic_audit.md",
        }
    ]

    converted = 0

    for markdown_path in markdown_files:
        relative = markdown_path.relative_to(
            SOURCE_ROOT
        )
        output_path = (
            DOCX_ROOT
            / relative.parent
            / f"{relative.stem}.docx"
        )

        create_docx(
            markdown_path,
            output_path,
        )
        converted += 1

    print("=" * 100)
    print("IIT JODHPUR — COMMAND 5 DOCX EXPORT")
    print("=" * 100)
    print()
    print("Markdown files:", len(markdown_files))
    print("DOCX files:", converted)
    print("Output:", DOCX_ROOT)


if __name__ == "__main__":
    main()