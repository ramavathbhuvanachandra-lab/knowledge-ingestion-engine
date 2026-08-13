"""
Phase 8.5 — Final RAG DOCX Builder

Input
-----
storage/organized_knowledge/<domain>/rag_plan.json

Output
------
storage/rag_knowledge/<domain>/*.docx
storage/rag_knowledge/<domain>/*.manifest.json
storage/rag_knowledge/<domain>/rag_build_manifest.json

Purpose
-------
Convert the validated 8.4 document plan into human-readable,
manually verifiable DOCX files.

Design
------
8.4 decides:
    WHAT belongs together.

8.5 decides:
    HOW that knowledge is presented inside the DOCX.

This phase does NOT:
- summarize
- rewrite factual knowledge
- call an LLM
- chunk
- embed
- create vectors
- perform retrieval
"""

from __future__ import annotations

from collections import OrderedDict
from pathlib import Path
import json
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class RAGDocumentBuilder:
    """
    Build final human-verifiable RAG DOCX files from rag_plan.json.
    """

    # ================================================================
    # INITIALIZATION
    # ================================================================

    def __init__(
        self,
        output_root: str | Path = (
            "storage/rag_knowledge"
        ),
    ):
        self.output_root = Path(
            output_root
        )

    # ================================================================
    # PUBLIC API
    # ================================================================

    def build_domain(
        self,
        organized_root: str | Path,
    ) -> dict:

        organized_root = Path(
            organized_root
        )

        if not organized_root.exists():
            raise FileNotFoundError(
                f"Organized knowledge root does not exist: "
                f"{organized_root}"
            )

        if not organized_root.is_dir():
            raise ValueError(
                f"Organized knowledge root is not a directory: "
                f"{organized_root}"
            )

        domain = organized_root.name

        # ------------------------------------------------------------
        # LOAD PLAN
        # ------------------------------------------------------------

        plan_path = (
            organized_root
            / "rag_plan.json"
        )

        if not plan_path.exists():
            raise FileNotFoundError(
                f"RAG plan does not exist: "
                f"{plan_path}"
            )

        plan = self._load_plan(
            plan_path
        )

        if plan.get("domain") != domain:
            raise ValueError(
                "RAG plan domain does not match "
                f"organized domain: "
                f"{plan.get('domain')} != {domain}"
            )

        documents = plan.get(
            "documents",
            [],
        )

        if not isinstance(
            documents,
            list,
        ):
            raise ValueError(
                "'documents' in RAG plan must be a list."
            )

        # ------------------------------------------------------------
        # OUTPUT ROOT
        # ------------------------------------------------------------

        output_root = (
            self.output_root
            / domain
        )

        output_root.mkdir(
            parents=True,
            exist_ok=True,
        )

        # Remove stale output from previous 8.5 runs.
        self._remove_existing_outputs(
            output_root
        )

        # ------------------------------------------------------------
        # INPUT INVENTORY
        # ------------------------------------------------------------

        input_unit_ids = (
            self._collect_plan_unit_ids(
                plan
            )
        )

        if not input_unit_ids:
            raise ValueError(
                "RAG plan contains no knowledge units."
            )

        # ------------------------------------------------------------
        # BUILD DOCUMENTS
        # ------------------------------------------------------------

        document_results = []

        output_unit_ids = []

        total_units = 0

        for document_plan in documents:

            if not isinstance(
                document_plan,
                dict,
            ):
                raise ValueError(
                    "Every RAG document plan must be an object."
                )

            document_id = self._clean_filename(
                document_plan.get(
                    "document_id"
                )
            )

            if not document_id:
                raise ValueError(
                    "RAG document is missing document_id."
                )

            title = (
                document_plan.get(
                    "title"
                )
                or self._display_name(
                    document_id
                )
            )

            units = document_plan.get(
                "units",
                [],
            )

            if not isinstance(
                units,
                list,
            ):
                raise ValueError(
                    f"Units must be a list for {document_id}."
                )

            if not units:
                continue

            output_path = (
                output_root
                / f"{document_id}.docx"
            )

            result = self._build_document(
                output_path=output_path,
                domain=domain,
                title=title,
                document_plan=document_plan,
            )

            document_results.append(
                result
            )

            current_ids = [
                unit["unit_id"]
                for unit in units
            ]

            output_unit_ids.extend(
                current_ids
            )

            total_units += len(
                current_ids
            )

        # ------------------------------------------------------------
        # OUTPUT COVERAGE
        # ------------------------------------------------------------

        duplicate_input_ids = (
            self._duplicates(
                input_unit_ids
            )
        )

        duplicate_output_ids = (
            self._duplicates(
                output_unit_ids
            )
        )

        missing_ids = sorted(
            set(input_unit_ids)
            - set(output_unit_ids)
        )

        extra_ids = sorted(
            set(output_unit_ids)
            - set(input_unit_ids)
        )

        if duplicate_input_ids:
            raise ValueError(
                "Duplicate input unit IDs: "
                f"{duplicate_input_ids}"
            )

        if duplicate_output_ids:
            raise ValueError(
                "Duplicate output unit IDs: "
                f"{duplicate_output_ids}"
            )

        if missing_ids:
            raise ValueError(
                "Knowledge units missing from final RAG output: "
                f"{missing_ids}"
            )

        if extra_ids:
            raise ValueError(
                "Unknown knowledge units found in final RAG output: "
                f"{extra_ids}"
            )

        if len(input_unit_ids) != len(
            output_unit_ids
        ):
            raise ValueError(
                "Final RAG unit count does not match "
                "planned unit count."
            )

        # ------------------------------------------------------------
        # GLOBAL BUILD MANIFEST
        # ------------------------------------------------------------

        global_manifest = {
            "phase": "8.5",

            "domain": domain,

            "input_plan": str(
                plan_path
            ),

            "output_root": str(
                output_root
            ),

            "documents_planned": len(
                documents
            ),

            "documents_built": len(
                document_results
            ),

            "input_units": len(
                input_unit_ids
            ),

            "output_units": len(
                output_unit_ids
            ),

            "coverage": {
                "input_units": len(
                    input_unit_ids
                ),

                "output_units": len(
                    output_unit_ids
                ),

                "missing_units": len(
                    missing_ids
                ),

                "extra_units": len(
                    extra_ids
                ),

                "duplicate_input_units": len(
                    duplicate_input_ids
                ),

                "duplicate_output_units": len(
                    duplicate_output_ids
                ),

                "coverage_complete": (
                    not missing_ids
                    and not extra_ids
                    and not duplicate_input_ids
                    and not duplicate_output_ids
                    and (
                        len(input_unit_ids)
                        == len(output_unit_ids)
                    )
                ),
            },

            "documents": document_results,
        }

        global_manifest_path = (
            output_root
            / "rag_build_manifest.json"
        )

        global_manifest_path.write_text(
            json.dumps(
                global_manifest,
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        # ------------------------------------------------------------
        # SUMMARY
        # ------------------------------------------------------------

        print()
        print("=" * 100)
        print("PHASE 8.5 — FINAL RAG DOCX BUILDER")
        print("=" * 100)

        print(
            "Domain              :",
            domain,
        )

        print(
            "DOCX Files Built     :",
            len(document_results),
        )

        print(
            "Input Units          :",
            len(input_unit_ids),
        )

        print(
            "Output Units         :",
            len(output_unit_ids),
        )

        print(
            "Coverage             :",
            (
                "PASS"
                if global_manifest[
                    "coverage"
                ][
                    "coverage_complete"
                ]
                else "FAIL"
            ),
        )

        print()
        print(
            "FINAL RAG FILES"
        )

        for result in document_results:

            print(
                f"{result['filename']:60}"
                f" → "
                f"{result['unit_count']:4} units"
            )

        print()
        print(
            "Build manifest:",
            global_manifest_path,
        )

        return global_manifest

    # ================================================================
    # BUILD ONE DOCUMENT
    # ================================================================

    def _build_document(
        self,
        *,
        output_path: Path,
        domain: str,
        title: str,
        document_plan: dict,
    ) -> dict:

        document = Document()

        self._configure_styles(
            document
        )

        units = document_plan.get(
            "units",
            [],
        )

        # ------------------------------------------------------------
        # TITLE
        # ------------------------------------------------------------

        title_heading = (
            document.add_heading(
                title,
                level=0,
            )
        )

        title_heading.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        # ------------------------------------------------------------
        # DOCUMENT OVERVIEW
        # ------------------------------------------------------------

        self._add_metadata_block(
            document,
            domain=domain,
            document_plan=document_plan,
        )

        document.add_paragraph()

        # ------------------------------------------------------------
        # GROUP UNITS INTO HUMAN-READABLE HIERARCHY
        # ------------------------------------------------------------

        grouped = (
            self._group_units_for_document(
                units
            )
        )

        # ------------------------------------------------------------
        # WRITE HIERARCHY
        # ------------------------------------------------------------

        for major_category, topics in (
            grouped.items()
        ):

            document.add_heading(
                major_category,
                level=1,
            )

            for topic, topic_units in (
                topics.items()
            ):

                document.add_heading(
                    topic,
                    level=2,
                )

                for unit in topic_units:

                    self._write_unit(
                        document,
                        unit,
                    )

        # ------------------------------------------------------------
        # SAVE
        # ------------------------------------------------------------

        document.save(
            output_path
        )

        # ------------------------------------------------------------
        # PER-DOCUMENT MANIFEST
        # ------------------------------------------------------------

        unit_ids = [
            unit["unit_id"]
            for unit in units
        ]

        source_documents = sorted(
            {
                unit[
                    "source_document"
                ]
                for unit in units
            }
        )

        source_urls = sorted(
            {
                unit[
                    "source_url"
                ]
                for unit in units
                if unit.get(
                    "source_url"
                )
            }
        )

        taxonomy_categories = sorted(
            {
                self._taxonomy_key(
                    unit
                )
                for unit in units
            }
        )

        manifest = {
            "phase": "8.5",

            "document_id": (
                document_plan.get(
                    "document_id"
                )
            ),

            "filename": output_path.name,

            "title": title,

            "domain": domain,

            "unit_count": len(
                units
            ),

            "unit_ids": unit_ids,

            "source_documents": (
                source_documents
            ),

            "source_urls": (
                source_urls
            ),

            "taxonomy_categories": (
                taxonomy_categories
            ),

            "coverage": {
                "unit_count_matches_plan": (
                    len(units)
                    == document_plan.get(
                        "unit_count",
                        len(units),
                    )
                ),
            },
        }

        manifest_path = (
            output_path.parent
            / (
                f"{output_path.stem}"
                ".manifest.json"
            )
        )

        manifest_path.write_text(
            json.dumps(
                manifest,
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        return {
            "document_id": (
                document_plan.get(
                    "document_id"
                )
            ),

            "filename": output_path.name,

            "path": str(
                output_path
            ),

            "manifest_path": str(
                manifest_path
            ),

            "title": title,

            "unit_count": len(
                units
            ),

            "source_documents": (
                len(source_documents)
            ),

            "source_urls": len(
                source_urls
            ),
        }

    # ================================================================
    # DOCUMENT GROUPING
    # ================================================================

    def _group_units_for_document(
        self,
        units: list[dict],
    ) -> OrderedDict:

        groups = OrderedDict()

        sorted_units = sorted(
            units,
            key=self._unit_sort_key,
        )

        for unit in sorted_units:

            major_category = (
                self._major_category(
                    unit
                )
            )

            topic = (
                self._specific_topic(
                    unit
                )
            )

            if major_category not in groups:
                groups[
                    major_category
                ] = OrderedDict()

            if topic not in groups[
                major_category
            ]:
                groups[
                    major_category
                ][topic] = []

            groups[
                major_category
            ][topic].append(
                unit
            )

        return groups

    # ================================================================
    # MAJOR CATEGORY
    # ================================================================

    def _major_category(
        self,
        unit: dict,
    ) -> str:

        taxonomy_domain = (
            unit.get(
                "taxonomy_domain"
            )
            or unit.get(
                "topic"
            )
            or "General"
        )

        taxonomy_category = (
            unit.get(
                "taxonomy_category"
            )
            or unit.get(
                "subtopic"
            )
            or ""
        )

        # Prefer actual section hierarchy when it contains a
        # meaningful parent heading.
        section_path = unit.get(
            "section_path",
            [],
        )

        if (
            isinstance(
                section_path,
                list,
            )
            and len(section_path) >= 1
        ):
            parent = str(
                section_path[0]
            ).strip()

            if (
                parent
                and len(parent) <= 120
            ):
                return parent

        # Otherwise use taxonomy category.
        if taxonomy_category:
            return self._display_name(
                taxonomy_category
            )

        return self._display_name(
            taxonomy_domain
        )

    # ================================================================
    # SPECIFIC TOPIC
    # ================================================================

    def _specific_topic(
        self,
        unit: dict,
    ) -> str:

        heading = (
            unit.get(
                "heading"
            )
            or ""
        ).strip()

        if heading:
            return self._clean_heading(
                heading
            )

        section_path = unit.get(
            "section_path",
            [],
        )

        if (
            isinstance(
                section_path,
                list,
            )
            and section_path
        ):

            return self._clean_heading(
                str(
                    section_path[-1]
                )
            )

        return "General Information"

    # ================================================================
    # WRITE UNIT
    # ================================================================

    def _write_unit(
        self,
        document: Document,
        unit: dict,
    ) -> None:

        heading = self._clean_heading(
            unit.get(
                "heading"
            )
            or "General Information"
        )

        # ------------------------------------------------------------
        # Specific topic
        # ------------------------------------------------------------

        document.add_heading(
            heading,
            level=3,
        )

        # ------------------------------------------------------------
        # Lightweight retrieval signals
        # ------------------------------------------------------------

        keywords = self._derive_keywords(
            unit
        )

        if keywords:

            paragraph = (
                document.add_paragraph()
            )

            run = paragraph.add_run(
                "Keywords: "
            )

            run.bold = True

            paragraph.add_run(
                ", ".join(
                    keywords
                )
            )

        # ------------------------------------------------------------
        # ORIGINAL KNOWLEDGE
        # ------------------------------------------------------------

        text = (
            unit.get(
                "text",
                "",
            )
            or ""
        ).strip()

        if text:

            self._add_original_content(
                document,
                text,
            )

        # ------------------------------------------------------------
        # SOURCE
        # ------------------------------------------------------------

        source_url = (
            unit.get(
                "source_url"
            )
            or ""
        ).strip()

        source_document = (
            unit.get(
                "source_document_name"
            )
            or unit.get(
                "source_document"
            )
            or ""
        ).strip()

        source_paragraph = (
            document.add_paragraph()
        )

        source_run = (
            source_paragraph.add_run(
                "Source: "
            )
        )

        source_run.bold = True

        if source_url:
            source_paragraph.add_run(
                source_url
            )

        elif source_document:
            source_paragraph.add_run(
                source_document
            )

        else:
            source_paragraph.add_run(
                "Unknown"
            )

        if (
            source_document
            and source_url
        ):

            source_document_paragraph = (
                document.add_paragraph()
            )

            document_run = (
                source_document_paragraph
                .add_run(
                    "Source Document: "
                )
            )

            document_run.bold = True

            source_document_paragraph.add_run(
                source_document
            )

        # ------------------------------------------------------------
        # OPTIONAL PAGE / SECTION TRACE
        # ------------------------------------------------------------

        page_number = unit.get(
            "page_number"
        )

        if page_number is not None:

            trace_paragraph = (
                document.add_paragraph()
            )

            trace_run = (
                trace_paragraph.add_run(
                    "Source Page: "
                )
            )

            trace_run.bold = True

            trace_paragraph.add_run(
                str(page_number)
            )

        # ------------------------------------------------------------
        # UNIT SEPARATOR
        # ------------------------------------------------------------

        separator = (
            document.add_paragraph()
        )

        separator.add_run(
            "—" * 30
        )

    # ================================================================
    # KEYWORDS
    # ================================================================

    def _derive_keywords(
        self,
        unit: dict,
    ) -> list[str]:

        candidates = []

        # Heading is the strongest retrieval signal.
        heading = (
            unit.get(
                "heading"
            )
            or ""
        ).strip()

        if heading:
            candidates.extend(
                self._extract_phrases(
                    heading
                )
            )

        # Taxonomy category.
        category = (
            unit.get(
                "taxonomy_category"
            )
            or ""
        ).strip()

        if category:
            candidates.append(
                self._display_name(
                    category
                )
            )

        # Subcategory.
        subcategory = (
            unit.get(
                "taxonomy_subcategory"
            )
            or ""
        ).strip()

        if (
            subcategory
            and subcategory != category
        ):
            candidates.append(
                self._display_name(
                    subcategory
                )
            )

        # First hierarchy parent can be useful.
        section_path = unit.get(
            "section_path",
            [],
        )

        if (
            isinstance(
                section_path,
                list,
            )
            and section_path
        ):
            parent = str(
                section_path[0]
            ).strip()

            if (
                parent
                and parent.lower()
                != heading.lower()
            ):
                candidates.append(
                    parent
                )

        return self._dedupe_keywords(
            candidates
        )[:6]

    # ================================================================
    # KEYWORD EXTRACTION
    # ================================================================

    def _extract_phrases(
        self,
        heading: str,
    ) -> list[str]:

        heading = self._clean_heading(
            heading
        )

        if not heading:
            return []

        candidates = [
            heading
        ]

        # Preserve common acronym notation.
        acronyms = re.findall(
            r"\b[A-Z][A-Z0-9.&-]{1,8}\b",
            heading,
        )

        candidates.extend(
            acronyms
        )

        # Parenthetical phrase.
        parentheses = re.findall(
            r"\(([^)]+)\)",
            heading,
        )

        candidates.extend(
            parentheses
        )

        return candidates

    # ================================================================
    # ORIGINAL CONTENT
    # ================================================================

    def _add_original_content(
        self,
        document: Document,
        text: str,
    ) -> None:

        # Keep original information intact while preserving obvious
        # paragraph boundaries.

        paragraphs = re.split(
            r"\n{2,}",
            text,
        )

        if len(paragraphs) == 1:

            lines = text.splitlines()

            for line in lines:

                line = line.strip()

                if line:
                    document.add_paragraph(
                        line
                    )

            return

        for paragraph_text in paragraphs:

            paragraph_text = (
                paragraph_text.strip()
            )

            if not paragraph_text:
                continue

            document.add_paragraph(
                paragraph_text
            )

    # ================================================================
    # METADATA BLOCK
    # ================================================================

    def _add_metadata_block(
        self,
        document: Document,
        *,
        domain: str,
        document_plan: dict,
    ) -> None:

        self._add_metadata_line(
            document.add_paragraph(),
            "Domain",
            domain,
        )

        self._add_metadata_line(
            document.add_paragraph(),
            "RAG Document",
            document_plan.get(
                "document_id",
                "",
            ),
        )

        unit_count = document_plan.get(
            "unit_count"
        )

        if unit_count is not None:

            self._add_metadata_line(
                document.add_paragraph(),
                "Knowledge Units",
                str(
                    unit_count
                ),
            )

        topics = document_plan.get(
            "topics",
            [],
        )

        if topics:

            self._add_metadata_line(
                document.add_paragraph(),
                "Topics",
                ", ".join(
                    topics
                ),
            )

        categories = document_plan.get(
            "taxonomy_categories",
            [],
        )

        if categories:

            self._add_metadata_line(
                document.add_paragraph(),
                "Taxonomy",
                "; ".join(
                    categories
                ),
            )

    # ================================================================
    # METADATA LINE
    # ================================================================

    def _add_metadata_line(
        self,
        paragraph,
        label: str,
        value: str,
    ) -> None:

        run = paragraph.add_run(
            f"{label}: "
        )

        run.bold = True

        paragraph.add_run(
            value or ""
        )

    # ================================================================
    # STYLE
    # ================================================================

    def _configure_styles(
        self,
        document: Document,
    ) -> None:

        styles = document.styles

        normal = styles[
            "Normal"
        ]

        normal.font.name = (
            "Arial"
        )

        normal.font.size = (
            Pt(10)
        )

        for style_name, size in (
            (
                "Heading 1",
                16,
            ),
            (
                "Heading 2",
                13,
            ),
            (
                "Heading 3",
                11,
            ),
        ):

            if style_name in styles:

                styles[
                    style_name
                ].font.name = (
                    "Arial"
                )

                styles[
                    style_name
                ].font.size = (
                    Pt(size)
                )

    # ================================================================
    # LOAD PLAN
    # ================================================================

    def _load_plan(
        self,
        path: Path,
    ) -> dict:

        try:

            data = json.loads(
                path.read_text(
                    encoding="utf-8"
                )
            )

        except json.JSONDecodeError as exc:

            raise ValueError(
                f"Invalid RAG plan JSON: "
                f"{path}"
            ) from exc

        if not isinstance(
            data,
            dict,
        ):
            raise ValueError(
                f"RAG plan must be a JSON object: "
                f"{path}"
            )

        return data

    # ================================================================
    # COLLECT PLAN IDS
    # ================================================================

    def _collect_plan_unit_ids(
        self,
        plan: dict,
    ) -> list[str]:

        ids = []

        for document in plan.get(
            "documents",
            [],
        ):

            for unit in document.get(
                "units",
                [],
            ):

                unit_id = unit.get(
                    "unit_id"
                )

                if unit_id:
                    ids.append(
                        unit_id
                    )

        return ids

    # ================================================================
    # CLEAN OUTPUT
    # ================================================================

    def _remove_existing_outputs(
        self,
        output_root: Path,
    ) -> None:

        for path in output_root.iterdir():

            if (
                path.is_file()
                and (
                    path.suffix.lower()
                    in {
                        ".docx",
                        ".json",
                    }
                )
            ):
                path.unlink()

    # ================================================================
    # CLEAN FILENAME
    # ================================================================

    def _clean_filename(
        self,
        value,
    ) -> str:

        if value is None:
            return ""

        value = str(
            value
        ).strip()

        value = re.sub(
            r"[^a-zA-Z0-9._-]+",
            "_",
            value,
        )

        return value.strip(
            "._"
        )

    # ================================================================
    # CLEAN HEADING
    # ================================================================

    def _clean_heading(
        self,
        value: str,
    ) -> str:

        value = (
            value
            or ""
        ).strip()

        # Remove Markdown link wrapper:
        #
        # [Professor Name](url)
        #
        # while preserving visible text.
        value = re.sub(
            r"\[([^\]]+)\]\([^)]+\)",
            r"\1",
            value,
        )

        # Remove leading/trailing markdown heading markers.
        value = re.sub(
            r"^#+\s*",
            "",
            value,
        )

        return value.strip()

    # ================================================================
    # DISPLAY NAME
    # ================================================================

    def _display_name(
        self,
        value: str,
    ) -> str:

        value = re.sub(
            r"[_\-]+",
            " ",
            value,
        )

        value = re.sub(
            r"\s+",
            " ",
            value,
        )

        return (
            value
            .strip()
            .title()
        )

    # ================================================================
    # TAXONOMY KEY
    # ================================================================

    def _taxonomy_key(
        self,
        unit: dict,
    ) -> str:

        domain = (
            unit.get(
                "taxonomy_domain"
            )
            or unit.get(
                "topic"
            )
            or "other"
        )

        category = (
            unit.get(
                "taxonomy_category"
            )
            or unit.get(
                "subtopic"
            )
            or "other"
        )

        subcategory = (
            unit.get(
                "taxonomy_subcategory"
            )
            or category
        )

        return (
            f"{domain}/"
            f"{category}/"
            f"{subcategory}"
        )

    # ================================================================
    # SORT
    # ================================================================

    def _unit_sort_key(
        self,
        unit: dict,
    ) -> tuple:

        try:
            page_index = int(
                unit.get(
                    "page_index",
                    0,
                )
            )
        except (
            TypeError,
            ValueError,
        ):
            page_index = 0

        try:
            section_index = int(
                unit.get(
                    "section_index",
                    0,
                )
            )
        except (
            TypeError,
            ValueError,
        ):
            section_index = 0

        return (
            unit.get(
                "source_document",
                "",
            ),
            page_index,
            section_index,
        )

    # ================================================================
    # DEDUPE KEYWORDS
    # ================================================================

    def _dedupe_keywords(
        self,
        values: list[str],
    ) -> list[str]:

        seen = set()
        result = []

        for value in values:

            cleaned = (
                value or ""
            ).strip()

            if not cleaned:
                continue

            key = cleaned.lower()

            if key in seen:
                continue

            seen.add(key)

            result.append(
                cleaned
            )

        return result

    # ================================================================
    # DUPLICATES
    # ================================================================

    def _duplicates(
        self,
        values: list[str],
    ) -> list[str]:

        seen = set()
        duplicates = set()

        for value in values:

            if value in seen:
                duplicates.add(
                    value
                )

            seen.add(
                value
            )

        return sorted(
            duplicates
        )


if __name__ == "__main__":

    print(
        "RAGDocumentBuilder"
    )

    print(
        "Phase 8.5 — Final RAG DOCX Builder"
    )
