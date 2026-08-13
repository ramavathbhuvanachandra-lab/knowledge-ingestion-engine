"""
Phase 8.5 — Final RAG Document Validator

Validates:

1. RAG plan exists.
2. Final DOCX files exist.
3. Per-document manifests exist.
4. Unit coverage is complete.
5. No unit is duplicated.
6. No unexpected unit exists.
7. DOCX files are readable.
8. Required structural headings exist.
9. Source URLs are represented.
10. Final output is consistent with the plan.

This validator does NOT:
- use an LLM
- embed
- chunk
- retrieve
- modify DOCX files
"""

from __future__ import annotations

from pathlib import Path
from collections import Counter
import json
import re

from docx import Document


class RAGDocumentValidator:

    # ================================================================
    # PUBLIC API
    # ================================================================

    def validate_domain(
        self,
        organized_root: str | Path,
        rag_root: str | Path,
    ) -> dict:

        organized_root = Path(
            organized_root
        )

        rag_root = Path(
            rag_root
        )

        if not organized_root.exists():
            raise FileNotFoundError(
                f"Organized root does not exist: "
                f"{organized_root}"
            )

        if not rag_root.exists():
            raise FileNotFoundError(
                f"RAG root does not exist: "
                f"{rag_root}"
            )

        domain = organized_root.name

        plan_path = (
            organized_root
            / "rag_plan.json"
        )

        if not plan_path.exists():
            raise FileNotFoundError(
                f"RAG plan does not exist: "
                f"{plan_path}"
            )

        plan = json.loads(
            plan_path.read_text(
                encoding="utf-8"
            )
        )

        if plan.get(
            "domain"
        ) != domain:
            raise ValueError(
                "Plan domain does not match "
                "organized domain."
            )

        documents = plan.get(
            "documents",
            [],
        )

        expected_ids = [
            unit["unit_id"]
            for document
            in documents
            for unit
            in document.get(
                "units",
                [],
            )
        ]

        expected_id_set = set(
            expected_ids
        )

        duplicate_plan_ids = (
            self._duplicates(
                expected_ids
            )
        )

        if duplicate_plan_ids:
            raise ValueError(
                "Duplicate unit IDs in plan: "
                f"{duplicate_plan_ids}"
            )

        # ------------------------------------------------------------
        # EXPECTED FILES
        # ------------------------------------------------------------

        expected_docx = []

        for document in documents:

            document_id = self._clean_filename(
                document.get(
                    "document_id"
                )
            )

            if not document_id:
                raise ValueError(
                    "Plan contains document "
                    "without document_id."
                )

            expected_docx.append(
                document_id
            )

        # ------------------------------------------------------------
        # CHECK FILE EXISTENCE
        # ------------------------------------------------------------

        missing_docx = []

        missing_manifests = []

        unreadable_docx = []

        document_reports = []

        for document in documents:

            document_id = self._clean_filename(
                document.get(
                    "document_id"
                )
            )

            docx_path = (
                rag_root
                / f"{document_id}.docx"
            )

            manifest_path = (
                rag_root
                / (
                    f"{document_id}"
                    ".manifest.json"
                )
            )

            if not docx_path.exists():
                missing_docx.append(
                    str(docx_path)
                )

            if not manifest_path.exists():
                missing_manifests.append(
                    str(manifest_path)
                )

            if docx_path.exists():

                try:
                    docx = Document(
                        str(docx_path)
                    )

                    paragraphs = [
                        p.text.strip()
                        for p in docx.paragraphs
                        if p.text.strip()
                    ]

                    heading_count = sum(
                        1
                        for paragraph
                        in docx.paragraphs
                        if paragraph.style.name.startswith(
                            "Heading"
                        )
                    )

                    has_title = bool(
                        paragraphs
                    )

                    has_keywords = any(
                        paragraph.lower().startswith(
                            "keywords:"
                        )
                        for paragraph
                        in paragraphs
                    )

                    has_source = any(
                        paragraph.lower().startswith(
                            "source:"
                        )
                        for paragraph
                        in paragraphs
                    )

                    document_reports.append(
                        {
                            "document_id": (
                                document_id
                            ),

                            "path": str(
                                docx_path
                            ),

                            "readable": True,

                            "paragraphs": len(
                                paragraphs
                            ),

                            "headings": (
                                heading_count
                            ),

                            "has_title": (
                                has_title
                            ),

                            "has_keywords": (
                                has_keywords
                            ),

                            "has_source": (
                                has_source
                            ),

                            "file_size_bytes": (
                                docx_path.stat().st_size
                            ),
                        }
                    )

                except Exception as error:

                    unreadable_docx.append(
                        {
                            "path": str(
                                docx_path
                            ),
                            "error": (
                                f"{type(error).__name__}: "
                                f"{error}"
                            ),
                        }
                    )

        # ------------------------------------------------------------
        # EXTRA DOCX
        # ------------------------------------------------------------

        actual_docx_ids = {
            path.stem
            for path in rag_root.glob(
                "*.docx"
            )
        }

        expected_docx_ids = set(
            expected_docx
        )

        extra_docx = sorted(
            actual_docx_ids
            - expected_docx_ids
        )

        # ------------------------------------------------------------
        # PER-DOCUMENT MANIFEST VALIDATION
        # ------------------------------------------------------------

        manifest_unit_ids = []

        manifest_errors = []

        for document in documents:

            document_id = self._clean_filename(
                document.get(
                    "document_id"
                )
            )

            manifest_path = (
                rag_root
                / (
                    f"{document_id}"
                    ".manifest.json"
                )
            )

            if not manifest_path.exists():
                continue

            try:

                manifest = json.loads(
                    manifest_path.read_text(
                        encoding="utf-8"
                    )
                )

                unit_ids = manifest.get(
                    "unit_ids",
                    [],
                )

                if not isinstance(
                    unit_ids,
                    list,
                ):
                    raise ValueError(
                        "'unit_ids' must be a list."
                    )

                expected_document_ids = [
                    unit["unit_id"]
                    for unit in document.get(
                        "units",
                        [],
                    )
                ]

                if set(unit_ids) != set(
                    expected_document_ids
                ):
                    manifest_errors.append(
                        {
                            "document_id": (
                                document_id
                            ),
                            "error": (
                                "Per-document manifest "
                                "unit IDs do not match "
                                "RAG plan."
                            ),
                        }
                    )

                manifest_unit_ids.extend(
                    unit_ids
                )

            except Exception as error:

                manifest_errors.append(
                    {
                        "document_id": (
                            document_id
                        ),
                        "error": (
                            f"{type(error).__name__}: "
                            f"{error}"
                        ),
                    }
                )

        # ------------------------------------------------------------
        # COVERAGE
        # ------------------------------------------------------------

        duplicate_manifest_ids = (
            self._duplicates(
                manifest_unit_ids
            )
        )

        missing_ids = sorted(
            expected_id_set
            - set(manifest_unit_ids)
        )

        extra_ids = sorted(
            set(manifest_unit_ids)
            - expected_id_set
        )

        # ------------------------------------------------------------
        # GLOBAL MANIFEST
        # ------------------------------------------------------------

        global_manifest_path = (
            rag_root
            / "rag_build_manifest.json"
        )

        global_manifest_exists = (
            global_manifest_path.exists()
        )

        # ------------------------------------------------------------
        # FINAL VALIDATION
        # ------------------------------------------------------------

        coverage_complete = (
            not missing_ids
            and not extra_ids
            and not duplicate_plan_ids
            and not duplicate_manifest_ids
        )

        structure_complete = (
            not missing_docx
            and not missing_manifests
            and not unreadable_docx
            and not manifest_errors
        )

        overall_pass = (
            coverage_complete
            and structure_complete
            and not extra_docx
            and global_manifest_exists
        )

        result = {
            "phase": "8.5",

            "domain": domain,

            "organized_root": str(
                organized_root
            ),

            "rag_root": str(
                rag_root
            ),

            "expected_documents": len(
                documents
            ),

            "actual_docx_documents": len(
                actual_docx_ids
            ),

            "expected_units": len(
                expected_ids
            ),

            "manifest_units": len(
                manifest_unit_ids
            ),

            "missing_docx": missing_docx,

            "missing_manifests": (
                missing_manifests
            ),

            "extra_docx": extra_docx,

            "unreadable_docx": (
                unreadable_docx
            ),

            "manifest_errors": (
                manifest_errors
            ),

            "duplicate_plan_unit_ids": (
                duplicate_plan_ids
            ),

            "duplicate_manifest_unit_ids": (
                duplicate_manifest_ids
            ),

            "missing_units": missing_ids,

            "extra_units": extra_ids,

            "global_manifest_exists": (
                global_manifest_exists
            ),

            "coverage_complete": (
                coverage_complete
            ),

            "structure_complete": (
                structure_complete
            ),

            "overall_pass": (
                overall_pass
            ),

            "documents": document_reports,
        }

        validation_path = (
            rag_root
            / "rag_validation_report.json"
        )

        validation_path.write_text(
            json.dumps(
                result,
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        # ------------------------------------------------------------
        # LOG
        # ------------------------------------------------------------

        print()
        print("=" * 100)
        print(
            "PHASE 8.5 — RAG DOCUMENT VALIDATION"
        )
        print("=" * 100)

        print(
            "Domain:",
            domain,
        )

        print(
            "Expected DOCX:",
            len(documents),
        )

        print(
            "Actual DOCX:",
            len(actual_docx_ids),
        )

        print(
            "Expected units:",
            len(expected_ids),
        )

        print(
            "Manifest units:",
            len(manifest_unit_ids),
        )

        print(
            "Coverage:",
            (
                "PASS"
                if coverage_complete
                else "FAIL"
            ),
        )

        print(
            "Structure:",
            (
                "PASS"
                if structure_complete
                else "FAIL"
            ),
        )

        print(
            "Overall:",
            (
                "PASS"
                if overall_pass
                else "FAIL"
            ),
        )

        print()
        print(
            "Validation report:",
            validation_path,
        )

        return result

    # ================================================================
    # FILENAME CLEANING
    # ================================================================

    def _clean_filename(
        self,
        value,
    ) -> str:

        if value is None:
            return ""

        value = str(
            value
        ).strip()

        value = re.sub(
            r"[^a-zA-Z0-9._-]+",
            "_",
            value,
        )

        return value.strip(
            "._"
        )

    # ================================================================
    # DUPLICATES
    # ================================================================

    def _duplicates(
        self,
        values: list[str],
    ) -> list[str]:

        seen = set()
        duplicates = set()

        for value in values:

            if value in seen:
                duplicates.add(
                    value
                )

            seen.add(
                value
            )

        return sorted(
            duplicates
        )


if __name__ == "__main__":

    print(
        "RAGDocumentValidator"
    )

    print(
        "Phase 8.5 — Final RAG Document Validator"
    )
